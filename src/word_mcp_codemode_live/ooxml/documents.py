"""Minimal closed-DOCX metadata helpers."""

from typing import Any

from docx import Document

_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_TEXT_TAG = f"{{{_WORD_NS}}}t"
_MOVE_FROM_TAG = f"{{{_WORD_NS}}}moveFrom"


def get_effective_text(paragraph: Any) -> str:
    """Return accepted-view paragraph text, including tracked insertions."""
    texts: list[str] = []
    for text_element in paragraph._element.iter(_TEXT_TAG):
        ancestor = text_element.getparent()
        while ancestor is not None and ancestor is not paragraph._element:
            if ancestor.tag == _MOVE_FROM_TAG:
                break
            ancestor = ancestor.getparent()
        else:
            if text_element.text:
                texts.append(text_element.text)
    return "".join(texts)


def get_document_properties(doc_path: str) -> dict[str, Any]:
    """Return cheap DOCX metadata without pretending to know rendered pages."""
    document = Document(doc_path)
    properties = document.core_properties
    return {
        "title": properties.title or "",
        "author": properties.author or "",
        "subject": properties.subject or "",
        "keywords": properties.keywords or "",
        "created": str(properties.created) if properties.created else "",
        "modified": str(properties.modified) if properties.modified else "",
        "last_modified_by": properties.last_modified_by or "",
        "revision": properties.revision or 0,
        "section_count": len(document.sections),
        "body_paragraph_word_count": sum(
            len(get_effective_text(paragraph).split()) for paragraph in document.paragraphs
        ),
        "body_paragraph_count": len(document.paragraphs),
        "body_table_count": len(document.tables),
    }
