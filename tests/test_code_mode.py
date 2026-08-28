import pytest
from docx import Document
from fastmcp import Client

from word_mcp_codemode_live.main import create_server


@pytest.mark.asyncio
async def test_code_mode_discovers_batch_and_page_capture_tools() -> None:
    async with Client(create_server(tool_mode="code")) as client:
        result = await client.call_tool(
            "search", {"query": "atomic batch edit verify screenshot Word pages"}
        )

    text = result.content[0].text
    assert "word_live_edit_batch" in text
    assert "word_live_capture_pages" in text


@pytest.mark.asyncio
async def test_code_mode_edits_and_merges_real_documents(tmp_path) -> None:
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    merged = tmp_path / "merged.docx"
    operations = [
        ("create_document", {"filename": str(first), "title": "First"}),
        ("add_heading", {"filename": str(first), "text": "Code Mode Report", "level": 1}),
        ("add_paragraph", {"filename": str(first), "text": "First document content."}),
        ("create_document", {"filename": str(second), "title": "Second"}),
        ("add_paragraph", {"filename": str(second), "text": "Second document content."}),
        (
            "merge_documents",
            {
                "target_filename": str(merged),
                "source_filenames": [str(first), str(second)],
                "add_page_breaks": True,
            },
        ),
    ]
    code = "\n".join(f"await call_tool({name!r}, {arguments!r})" for name, arguments in operations)
    code += f"\nreturn await call_tool('get_document_text', {{'filename': {str(merged)!r}}})"

    async with Client(create_server(tool_mode="code")) as client:
        result = await client.call_tool("execute", {"code": code})

    merged_document = Document(merged)
    merged_text = "\n".join(paragraph.text for paragraph in merged_document.paragraphs)
    assert not result.is_error
    assert "Code Mode Report" in merged_text
    assert "First document content." in merged_text
    assert "Second document content." in merged_text
