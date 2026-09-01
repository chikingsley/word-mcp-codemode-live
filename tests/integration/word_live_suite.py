"""Run every registered tool through FastMCP against disposable Word documents.

This is an explicit Windows integration runner, not part of the normal unit test
suite. It launches a private Word instance, creates all documents under a temporary
directory, calls tools through ``fastmcp.Client``, verifies Word/file state after
each call, writes reports, and removes the temporary directory.
"""

from __future__ import annotations

import argparse
import asyncio
import hashlib
import json
import logging
import platform
import subprocess
import sys
import tempfile
import time
from collections.abc import Callable
from dataclasses import asdict, dataclass, field
from importlib.metadata import version
from pathlib import Path
from typing import Any

import pymupdf
import win32com.client
from docx import Document as DocxDocument
from fastmcp import Client

from word_mcp_codemode_live.core import word_com
from word_mcp_codemode_live.main import create_server

logger = logging.getLogger(__name__)

ArgsFactory = Callable[[Any, Path], dict[str, Any]]
Setup = Callable[[Any, Any, Path], None]
Verifier = Callable[[Any, Any, dict[str, Any], Path], tuple[bool, str]]


@dataclass(frozen=True, slots=True)
class PreCall:
    name: str
    arguments: ArgsFactory


@dataclass(frozen=True, slots=True)
class LiveCase:
    name: str
    arguments: ArgsFactory
    verifier: Verifier
    setup: Setup | None = None
    pre_calls: tuple[PreCall, ...] = ()


@dataclass(slots=True)
class ResultRecord:
    tool: str
    status: str
    invocation: str
    elapsed_seconds: float
    mcp_is_error: bool
    nested_error: str | None
    postcondition: str
    response_excerpt: str
    arguments: dict[str, Any] = field(default_factory=dict)


@dataclass(slots=True)
class WorkflowRecord:
    workflow: str
    status: str
    elapsed_seconds: float
    steps_completed: int
    expected_steps: int
    evidence: str
    error: str | None = None


def _json_safe(value: Any) -> Any:
    try:
        json.dumps(value)
        return value
    except TypeError:
        return repr(value)


def _response_payload(result: Any) -> tuple[dict[str, Any], str, str | None]:
    text_parts = [
        str(item.text) for item in result.content if getattr(item, "type", None) == "text"
    ]
    response_text = "\n".join(text_parts)
    payload: Any = None
    if response_text:
        try:
            payload = json.loads(response_text)
        except json.JSONDecodeError:
            payload = None

    if not isinstance(payload, dict):
        structured = getattr(result, "structured_content", None)
        if isinstance(structured, dict):
            nested = structured.get("result")
            if isinstance(nested, str):
                try:
                    payload = json.loads(nested)
                except json.JSONDecodeError:
                    payload = structured
            else:
                payload = structured

    if not isinstance(payload, dict):
        payload = {"raw_text": response_text}

    nested_error = payload.get("error")
    if nested_error is None and payload.get("success") is False:
        nested_error = payload.get("message") or "success=false"
    return payload, response_text, str(nested_error) if nested_error is not None else None


async def _invoke(
    client: Client, name: str, arguments: dict[str, Any]
) -> tuple[Any, dict, str, str | None]:
    result = await client.call_tool(name, arguments)
    payload, response_text, nested_error = _response_payload(result)
    return result, payload, response_text, nested_error


def _make_closed_docx(path: Path) -> None:
    document = DocxDocument()
    document.add_heading("Capability verification", level=1)
    document.add_paragraph("Alpha marker")
    document.add_paragraph("Beta marker")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Seed A1"
    table.cell(0, 1).text = "Seed A2"
    table.cell(1, 0).text = "Seed B1"
    table.cell(1, 1).text = "Seed B2"
    document.save(str(path))


def _make_png(path: Path) -> None:
    pdf = pymupdf.open()
    page = pdf.new_page(width=240, height=120)
    page.insert_text((24, 64), "Word MCP verification")
    pixmap = page.get_pixmap(alpha=False)
    pixmap.save(str(path))
    pdf.close()


def _find_text_range(document: Any, text: str) -> Any:
    word_range = document.Content.Duplicate
    word_range.Find.ClearFormatting()
    if not word_range.Find.Execute(FindText=text, Forward=True, Wrap=0):
        raise RuntimeError(f"Seed text not found: {text}")
    return word_range.Duplicate


def _make_live_document(app: Any, path: Path) -> Any:
    document = app.Documents.Add()
    document.Content.Text = "Alpha marker\rBeta marker\rGamma marker\r"
    document.Paragraphs(1).Style = "Heading 1"

    table_range = document.Range(document.Content.End - 1, document.Content.End - 1)
    table = document.Tables.Add(table_range, 2, 2)
    table.Cell(1, 1).Range.Text = "Seed A1"
    table.Cell(1, 2).Range.Text = "Seed A2"
    table.Cell(2, 1).Range.Text = "Seed B1"
    table.Cell(2, 2).Range.Text = "Seed B2"

    comment_range = _find_text_range(document, "Alpha marker")
    document.Comments.Add(comment_range, "Seed comment")

    note_anchor = _find_text_range(document, "Beta marker")
    note_anchor.Collapse(0)
    note = document.Footnotes.Add(Range=note_anchor)
    note.Range.Text = "Seed footnote"

    document.Sections(1).Headers(1).Range.Text = "Seed header"
    document.Sections(1).Footers(1).Range.Text = "Seed footer"
    document.BuiltInDocumentProperties("Title").Value = "Seed title"
    document.SaveAs2(str(path), AddToRecentFiles=False)
    return document


def _add_revision(_app: Any, document: Any, _root: Path) -> None:
    document.TrackRevisions = True
    insertion = document.Range(document.Content.End - 1, document.Content.End - 1)
    insertion.InsertAfter("Revision marker")
    document.TrackRevisions = False


def _add_comment_reply(_app: Any, document: Any, _root: Path) -> None:
    root_comment = document.Comments(1)
    root_comment.Replies.Add(root_comment.Scope, "Seed reply")


def _add_date_field(_app: Any, document: Any, _root: Path) -> None:
    offset = int(document.Content.End) - 1
    document.Fields.Add(
        Range=document.Range(offset, offset),
        Type=-1,  # wdFieldEmpty
        Text="DATE",
        PreserveFormatting=True,
    )


def _add_unlinkable_fields(_app: Any, document: Any, _root: Path) -> None:
    _add_date_field(_app, document, _root)
    offset = int(document.Content.End) - 1
    sequence = document.Fields.Add(
        Range=document.Range(offset, offset),
        Type=-1,  # wdFieldEmpty
        Text="SEQ capability_item",
        PreserveFormatting=True,
    )
    sequence.Update()


def _add_fields_for_inspection(_app: Any, document: Any, _root: Path) -> None:
    _add_date_field(_app, document, _root)
    offset = int(document.Content.End) - 1
    document.Fields.Add(
        Range=document.Range(offset, offset),
        Type=-1,
        Text='XE "Capability Entry"',
        PreserveFormatting=True,
    )


def _add_hyperlink(_app: Any, document: Any, _root: Path) -> None:
    document.Hyperlinks.Add(
        Anchor=_find_text_range(document, "Beta marker"),
        Address="https://example.com/original",
    )


def _add_table_of_contents(_app: Any, document: Any, _root: Path) -> None:
    document.TablesOfContents.Add(
        Range=document.Range(0, 0),
        UseHeadingStyles=True,
        UpperHeadingLevel=1,
        LowerHeadingLevel=3,
        UseFields=False,
        TableID="",
        RightAlignPageNumbers=True,
        IncludePageNumbers=True,
        AddedStyles="",
        UseHyperlinks=True,
        HidePageNumbersInWeb=True,
        UseOutlineLevels=True,
    )


def _add_custom_style(_app: Any, document: Any, _root: Path) -> None:
    document.Styles.Add(Name="CapabilityStyle", Type=1)  # wdStyleTypeParagraph


def _add_custom_character_style(_app: Any, document: Any, _root: Path) -> None:
    document.Styles.Add(Name="CapabilityCharacterStyle", Type=2)  # wdStyleTypeCharacter


def _add_highlight(_app: Any, document: Any, _root: Path) -> None:
    _find_text_range(document, "Beta marker").HighlightColorIndex = 7  # wdYellow


def _set_top_gutter(_app: Any, document: Any, _root: Path) -> None:
    setup = document.Sections(1).PageSetup
    setup.Gutter = 36
    setup.GutterPos = 1  # wdGutterPosTop


def _setup_layout_diagnostics(_app: Any, document: Any, _root: Path) -> None:
    document.Range(0, 0).InsertBefore("😀")
    section_offset = int(_find_text_range(document, "Beta marker").Start)
    document.Range(section_offset, section_offset).InsertBreak(Type=2)  # next-page section
    page_offset = int(_find_text_range(document, "Gamma marker").Start)
    document.Variables.Add(Name="CapabilityExpectedManualBreakOffset", Value=str(page_offset))
    document.Range(page_offset, page_offset).InsertBreak(Type=7)  # manual page break
    _set_top_gutter(_app, document, _root)


def _apply_direct_heading_numbering(_app: Any, document: Any, _root: Path) -> None:
    document.Paragraphs(1).Range.ListFormat.ApplyNumberDefault()


def _make_merge_source(app: Any, destination: Any, root: Path) -> None:
    source_path = root / "merge-source.docx"
    source = app.Documents.Add()
    try:
        source.Content.Text = "Merged source marker\rMerged field marker\rMerged notes marker\r"
        merged_style = source.Styles.Add(Name="CapabilityMergedStyle", Type=1)
        merged_style.Font.Italic = True
        source.Paragraphs(1).Style = merged_style

        field_anchor = _find_text_range(source, "Merged field marker")
        field_anchor.Collapse(0)
        source.Fields.Add(
            Range=field_anchor,
            Type=-1,
            Text="DATE",
            PreserveFormatting=True,
        )
        source.Comments.Add(_find_text_range(source, "Merged source marker"), "Merged comment")

        note_anchor = _find_text_range(source, "Merged notes marker")
        note_anchor.Collapse(1)
        source.Footnotes.Add(Range=note_anchor).Range.Text = "Merged footnote"
        endnote_anchor = _find_text_range(source, "Merged notes marker")
        endnote_anchor.Collapse(0)
        source.Endnotes.Add(Range=endnote_anchor).Range.Text = "Merged endnote"

        section_offset = int(source.Content.End) - 1
        source.Range(section_offset, section_offset).InsertBreak(Type=2)
        source.TrackRevisions = True
        revision_anchor = source.Range(int(source.Content.End) - 1, int(source.Content.End) - 1)
        revision_anchor.InsertAfter("Merged tracked revision")
        source.TrackRevisions = False
        source.SaveAs2(str(source_path), AddToRecentFiles=False)
    finally:
        source.Close(SaveChanges=False)
        destination.Activate()


def _verify_native_merge(
    _app: Any, document: Any, payload: dict[str, Any], _root: Path
) -> tuple[bool, str]:
    deltas = payload.get("deltas", {})
    text = str(document.Content.Text)
    preserved = {
        "sections": deltas.get("sections", 0) >= 1,
        "styles": deltas.get("styles", 0) >= 1 and _style_exists(document, "CapabilityMergedStyle"),
        "fields": deltas.get("fields", 0) >= 1,
        "comments": deltas.get("comments", 0) >= 1,
        "footnotes": deltas.get("footnotes", 0) >= 1,
        "endnotes": deltas.get("endnotes", 0) >= 1,
        "revisions": deltas.get("revisions", 0) >= 1,
    }
    inserted = payload.get("inserted_range", {})
    ok = (
        payload.get("success") is True
        and payload.get("native_operation") == "Range.InsertFile"
        and inserted.get("end", 0) > inserted.get("start", 0)
        and "Merged source marker" in text
        and "Merged tracked revision" in text
        and all(preserved.values())
    )
    return ok, f"preserved={preserved!r}; inserted_range={inserted!r}"


def _setup_navigation(_app: Any, document: Any, _root: Path) -> None:
    offset = int(_find_text_range(document, "Gamma marker").Start)
    document.Range(offset, offset).InsertBreak(Type=7)
    document.Save()


def _setup_saved_document_with_linked_header(_app: Any, document: Any, _root: Path) -> None:
    document.Sections(1).Headers(1).Range.Text = "Inherited capability header"
    section_offset = int(document.Content.End) - 1
    document.Range(section_offset, section_offset).InsertBreak(Type=2)
    document.Sections(2).Headers(1).LinkToPrevious = True
    document.Save()


def _verify_snapshot_capture(
    _app: Any, document: Any, payload: dict[str, Any], root: Path
) -> tuple[bool, str]:
    snapshot_path = root / "snapshot.json"
    snapshot = (
        json.loads(snapshot_path.read_text(encoding="utf-8")) if snapshot_path.is_file() else {}
    )
    header_rows = {
        (row.get("section_index"), row.get("kind"), row.get("variant")): row
        for row in snapshot.get("content", {}).get("headers_footers", [])
    }
    first = header_rows.get((1, "header", "primary"), {})
    second = header_rows.get((2, "header", "primary"), {})
    ok = (
        payload.get("success") is True
        and snapshot_path.is_file()
        and snapshot.get("content_sha256") == payload.get("content_sha256")
        and snapshot.get("envelope_sha256") == payload.get("envelope_sha256")
        and bool(document.Saved) is True
        and first.get("text") == "Inherited capability header"
        and second.get("linked_to_previous") is True
        and second.get("text") == "Inherited capability header"
    )
    return (
        ok,
        f"snapshot exists={snapshot_path.is_file()}; Word Saved={bool(document.Saved)!r}; "
        f"section headers={[first.get('text'), second.get('text')]!r}",
    )


def _verify_navigation(
    app: Any, document: Any, payload: dict[str, Any], _root: Path
) -> tuple[bool, str]:
    selection = payload.get("selection", {})
    ok = (
        payload.get("success") is True
        and payload.get("requested") == {"kind": "page", "page": 2}
        and selection.get("collapsed") is True
        and selection.get("active_end_page") == 2
        and str(app.ActiveDocument.FullName).casefold() == str(document.FullName).casefold()
        and bool(document.Saved) is True
        and payload.get("saved_before") is True
        and payload.get("saved_after") is True
        and payload.get("saved_state_unchanged") is True
    )
    return (
        ok,
        f"selection={selection!r}; Word Saved={bool(document.Saved)!r}; "
        f"visible={bool(app.Visible)!r}",
    )


def _style_exists(document: Any, style_name: str) -> bool:
    return any(
        str(document.Styles(index).NameLocal) == style_name
        for index in range(1, int(document.Styles.Count) + 1)
    )


def _verify_layout_diagnostics(
    _app: Any, document: Any, payload: dict[str, Any], _root: Path
) -> tuple[bool, str]:
    sections = payload.get("sections", [])
    manual_offsets = payload.get("manual_page_break_offsets", [])
    section_break_offsets = {
        int(document.Sections(index).Range.End) - 1
        for index in range(1, int(document.Sections.Count))
    }
    expected_manual_offsets = [int(document.Variables("CapabilityExpectedManualBreakOffset").Value)]
    second_section = document.Sections(2).Range
    second_start = second_section.Duplicate
    second_start.SetRange(int(second_section.Start), int(second_section.Start))
    expected_start_page = int(second_start.Information(3))
    full_section_end_page = int(second_section.Information(3))
    expected_usable_height = (
        float(document.Sections(1).PageSetup.PageHeight)
        - float(document.Sections(1).PageSetup.TopMargin)
        - float(document.Sections(1).PageSetup.BottomMargin)
        - 36
    )
    ok = (
        payload.get("success") is True
        and payload.get("page_count") == int(document.ComputeStatistics(2))
        and payload.get("section_count") == 2
        and len(sections) == 2
        and sections[0].get("margins_points", {}).get("gutter_position") == "top"
        and abs(sections[0].get("usable_height_points", 0) - expected_usable_height) < 0.1
        and payload.get("manual_page_break_count") == 1
        and manual_offsets == expected_manual_offsets
        and not (set(manual_offsets) & section_break_offsets)
        and sections[1].get("start_page") == expected_start_page
        and full_section_end_page > expected_start_page
    )
    return (
        ok,
        (
            f"pages={payload.get('page_count')!r}; sections={payload.get('section_count')!r}; "
            f"manual_offsets={manual_offsets!r}; expected_manual_offsets={expected_manual_offsets!r}; "
            f"section_offsets={sorted(section_break_offsets)!r}; "
            f"section2 start/end pages={expected_start_page}/{full_section_end_page}"
        ),
    )


def _ok_payload(_app: Any, _doc: Any, payload: dict[str, Any], _root: Path) -> tuple[bool, str]:
    ok = payload.get("success") is True
    return ok, f"payload success={payload.get('success')!r}"


def _contains_payload(key: str, expected: Any) -> Verifier:
    def verify(_app: Any, _doc: Any, payload: dict[str, Any], _root: Path) -> tuple[bool, str]:
        actual = payload.get(key)
        return actual == expected, f"payload {key}={actual!r}; expected {expected!r}"

    return verify


def _text_contains(expected: str) -> Verifier:
    def verify(_app: Any, document: Any, _payload: dict[str, Any], _root: Path) -> tuple[bool, str]:
        text = str(document.Content.Text)
        return expected in text, f"document contains {expected!r}={expected in text}"

    return verify


def _text_absent(expected: str) -> Verifier:
    def verify(_app: Any, document: Any, _payload: dict[str, Any], _root: Path) -> tuple[bool, str]:
        text = str(document.Content.Text)
        return expected not in text, f"document excludes {expected!r}={expected not in text}"

    return verify


def _live_cases(image_path: Path) -> list[LiveCase]:
    filename = lambda _doc, root: {"filename": str(root / "case.docx")}  # noqa: E731
    return [
        LiveCase(
            "word_live_list_open",
            lambda _doc, _root: {},
            lambda _app, _doc, payload, _root: (
                payload.get("success") is True and payload.get("count", 0) >= 1,
                f"reported open documents={payload.get('count')!r}",
            ),
        ),
        LiveCase(
            "word_live_get_text",
            filename,
            lambda _app, _doc, payload, _root: (
                payload.get("success") is True
                and any(
                    "Alpha marker" in item.get("text", "") for item in payload.get("paragraphs", [])
                ),
                "Alpha marker present in returned paragraphs",
            ),
        ),
        LiveCase("word_live_get_info", filename, _ok_payload),
        LiveCase(
            "word_live_find_text",
            lambda _doc, root: {"filename": str(root / "case.docx"), "search_text": "Beta marker"},
            lambda _app, _doc, payload, _root: (
                payload.get("success") is True and payload.get("match_count", 0) >= 1,
                f"matches={payload.get('match_count')!r}",
            ),
        ),
        LiveCase(
            "word_live_get_page_text",
            lambda _doc, root: {"filename": str(root / "case.docx"), "page": 1},
            lambda _app, _doc, payload, _root: (
                payload.get("success") is True and "Alpha marker" in json.dumps(payload),
                "page result contains Alpha marker",
            ),
        ),
        LiveCase(
            "word_live_get_paragraph_format",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "start_paragraph": 1,
                "end_paragraph": 1,
                "include_runs": True,
            },
            _ok_payload,
        ),
        LiveCase(
            "word_live_inspect_document_outline",
            filename,
            lambda _app, _document, payload, _root: (
                payload.get("success") is True
                and payload.get("outline_entry_count", 0) >= 1
                and any(entry.get("outline_level") == 1 for entry in payload.get("entries", [])),
                f"outline entries={payload.get('outline_entry_count')!r}",
            ),
        ),
        LiveCase(
            "word_live_inspect_highlighted_text",
            filename,
            lambda _app, _document, payload, _root: (
                payload.get("success") is True
                and any(
                    "Beta marker" in item.get("text", "") and item.get("color") == "yellow"
                    for item in payload.get("highlights", [])
                ),
                f"highlighted ranges={payload.get('highlight_count')!r}",
            ),
            setup=_add_highlight,
        ),
        LiveCase(
            "word_live_inspect_heading_numbering",
            filename,
            lambda _app, _document, payload, _root: (
                payload.get("success") is True
                and payload.get("heading_count", 0) >= 1
                and len(payload.get("heading_styles", [])) == 9,
                (
                    f"headings={payload.get('heading_count')!r}; "
                    f"heading styles={len(payload.get('heading_styles', []))}"
                ),
            ),
        ),
        LiveCase(
            "word_live_setup_heading_numbering",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "replace_existing": True,
                "number_formats": {"1": "Article %1"},
            },
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and payload.get("numbered_heading_count", 0) >= 1
                and payload.get("heading_styles", [{}])[0]
                .get("level_definition", {})
                .get("number_format")
                == "Article %1"
                and int(document.Paragraphs(1).Range.ListFormat.ListType) != 0,
                (
                    f"numbered headings={payload.get('numbered_heading_count')!r}; "
                    f"Heading 1 format={payload.get('heading_styles', [{}])[0].get('level_definition', {}).get('number_format')!r}; "
                    f"Word list type={document.Paragraphs(1).Range.ListFormat.ListType}"
                ),
            ),
            setup=_apply_direct_heading_numbering,
        ),
        LiveCase(
            "word_live_get_comments",
            filename,
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and payload.get("comment_count") == 1
                and payload.get("raw_comment_count") == int(document.Comments.Count)
                and len(payload.get("comments", [])[0].get("replies", [])) == 1,
                (
                    f"payload threads={payload.get('comment_count')!r}; "
                    f"raw Word comments={document.Comments.Count}"
                ),
            ),
            setup=_add_comment_reply,
        ),
        LiveCase(
            "word_live_set_comment_status",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "comment_index": 1,
                "resolved": True,
            },
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and payload.get("resolved") is True
                and bool(document.Comments(1).Done),
                (
                    f"payload resolved={payload.get('resolved')!r}; "
                    f"Word Done={bool(document.Comments(1).Done)}"
                ),
            ),
        ),
        LiveCase(
            "word_live_list_revisions",
            filename,
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and payload.get("revision_count") == int(document.Revisions.Count)
                and int(document.Revisions.Count) > 0,
                f"payload revisions={payload.get('revision_count')!r}; Word revisions={document.Revisions.Count}",
            ),
            setup=_add_revision,
        ),
        LiveCase(
            "word_live_list_footnotes_endnotes",
            lambda _doc, root: {"filename": str(root / "case.docx"), "note_type": "all"},
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and payload.get("footnote_count") == int(document.Footnotes.Count)
                and int(document.Footnotes.Count) >= 1,
                f"payload footnotes={payload.get('footnote_count')!r}; Word footnotes={document.Footnotes.Count}",
            ),
        ),
        LiveCase(
            "word_live_get_note_configuration",
            filename,
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and payload.get("footnotes", {}).get("count") == int(document.Footnotes.Count)
                and payload.get("endnotes", {}).get("count") == int(document.Endnotes.Count),
                (
                    f"configured footnotes={payload.get('footnotes', {}).get('count')!r}; "
                    f"Word footnotes={document.Footnotes.Count}"
                ),
            ),
        ),
        LiveCase(
            "word_live_get_headers_footers",
            lambda _doc, root: {"filename": str(root / "case.docx"), "section": "all"},
            lambda _app, _doc, payload, _root: (
                payload.get("success") is True
                and "Seed header" in json.dumps(payload)
                and "Seed footer" in json.dumps(payload),
                "returned native header and footer text",
            ),
        ),
        LiveCase(
            "word_live_get_undo_history",
            filename,
            _ok_payload,
            pre_calls=(
                PreCall(
                    "word_live_insert_text",
                    lambda _doc, root: {
                        "filename": str(root / "case.docx"),
                        "text": "Undo history marker",
                        "position": "end",
                    },
                ),
            ),
        ),
        LiveCase(
            "word_live_insert_text",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "text": "Inserted marker",
                "position": "end",
            },
            _text_contains("Inserted marker"),
        ),
        LiveCase(
            "word_live_replace_text",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "find_text": "Alpha marker",
                "replace_text": "Alpha replaced",
            },
            _text_contains("Alpha replaced"),
        ),
        LiveCase(
            "word_live_insert_paragraphs",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "paragraphs": ["Inserted paragraph"],
                "target_paragraph_index": 1,
                "position": "after",
            },
            _text_contains("Inserted paragraph"),
        ),
        LiveCase(
            "word_live_insert_page_break",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "position": "end",
            },
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and int(document.ComputeStatistics(2)) >= 2
                and "\x0c" in str(document.Content.Text),
                (
                    f"Word pages={document.ComputeStatistics(2)}; "
                    f"page break present={chr(12) in str(document.Content.Text)}"
                ),
            ),
        ),
        LiveCase(
            "word_live_delete_text",
            lambda doc, root: {
                "filename": str(root / "case.docx"),
                "start": int(_find_text_range(doc, "Gamma marker").Start),
                "end": int(_find_text_range(doc, "Gamma marker").End),
            },
            _text_absent("Gamma marker"),
        ),
        LiveCase(
            "word_live_format_text",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "start_paragraph": 1,
                "end_paragraph": 1,
                "bold": True,
                "font_size": 15,
            },
            lambda _app, document, _payload, _root: (
                bool(document.Paragraphs(1).Range.Font.Bold)
                and abs(float(document.Paragraphs(1).Range.Font.Size) - 15) < 0.1,
                f"Word bold={document.Paragraphs(1).Range.Font.Bold}; size={document.Paragraphs(1).Range.Font.Size}",
            ),
        ),
        LiveCase(
            "word_live_list_custom_styles",
            filename,
            lambda _app, _document, payload, _root: (
                payload.get("success") is True
                and any(
                    style.get("name") == "CapabilityStyle" for style in payload.get("styles", [])
                ),
                f"custom styles={payload.get('custom_style_count')!r}",
            ),
            setup=_add_custom_style,
        ),
        LiveCase(
            "word_live_create_custom_style",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "style_name": "CreatedCapabilityCharacterStyle",
                "style_type": "character",
                "font_size": 13,
                "bold": True,
            },
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and _style_exists(document, "CreatedCapabilityCharacterStyle")
                and int(document.Styles("CreatedCapabilityCharacterStyle").Type) == 2
                and bool(document.Styles("CreatedCapabilityCharacterStyle").Font.Bold),
                "Word custom character style was created with bold formatting",
            ),
        ),
        LiveCase(
            "word_live_update_custom_style",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "style_name": "CapabilityCharacterStyle",
                "font_size": 14,
                "italic": True,
            },
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and abs(float(document.Styles("CapabilityCharacterStyle").Font.Size) - 14) < 0.1
                and bool(document.Styles("CapabilityCharacterStyle").Font.Italic),
                (
                    f"Word style size={document.Styles('CapabilityCharacterStyle').Font.Size}; "
                    f"italic={document.Styles('CapabilityCharacterStyle').Font.Italic}"
                ),
            ),
            setup=_add_custom_character_style,
        ),
        LiveCase(
            "word_live_delete_custom_style",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "style_name": "CapabilityCharacterStyle",
            },
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and not _style_exists(document, "CapabilityCharacterStyle"),
                f"Word style still exists={_style_exists(document, 'CapabilityCharacterStyle')}",
            ),
            setup=_add_custom_character_style,
        ),
        LiveCase(
            "word_live_apply_list",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "start_paragraph": 2,
                "end_paragraph": 3,
                "list_type": "bullet",
            },
            lambda _app, document, _payload, _root: (
                int(document.Paragraphs(2).Range.ListFormat.ListType) != 0,
                f"Word list type={document.Paragraphs(2).Range.ListFormat.ListType}",
            ),
        ),
        LiveCase(
            "word_live_set_paragraph_spacing",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "paragraph_index": 2,
                "space_after_pt": 18,
                "alignment": "center",
            },
            lambda _app, document, _payload, _root: (
                abs(float(document.Paragraphs(2).Format.SpaceAfter) - 18) < 0.1
                and int(document.Paragraphs(2).Format.Alignment) == 1,
                f"Word space_after={document.Paragraphs(2).Format.SpaceAfter}; alignment={document.Paragraphs(2).Format.Alignment}",
            ),
        ),
        LiveCase(
            "word_live_add_table",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "rows": 2,
                "cols": 2,
                "data": [["New A1", "New A2"], ["New B1", "New B2"]],
                "position": "end",
            },
            lambda _app, document, _payload, _root: (
                int(document.Tables.Count) == 2
                and "New A1" in str(document.Tables(2).Cell(1, 1).Range.Text),
                f"Word table count={document.Tables.Count}",
            ),
        ),
        LiveCase(
            "word_live_format_table",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "table_index": 1,
                "table_alignment": "center",
                "font_size": 10,
                "cell_shading": [[1, 1, "D9EAF7"]],
            },
            lambda _app, document, _payload, _root: (
                int(document.Tables(1).Rows.Alignment) == 1
                and abs(float(document.Tables(1).Range.Font.Size) - 10) < 0.1,
                f"Word table alignment={document.Tables(1).Rows.Alignment}; font_size={document.Tables(1).Range.Font.Size}",
            ),
        ),
        LiveCase(
            "word_live_modify_table",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "table_index": 1,
                "operation": "set_cell",
                "row": 1,
                "col": 1,
                "text": "Changed cell",
            },
            lambda _app, document, _payload, _root: (
                "Changed cell" in str(document.Tables(1).Cell(1, 1).Range.Text),
                f"Word cell text={document.Tables(1).Cell(1, 1).Range.Text!r}",
            ),
        ),
        LiveCase(
            "word_live_add_comment",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "paragraph_index": 2,
                "text": "Added comment",
                "author": "Capability Verifier",
            },
            lambda _app, document, _payload, _root: (
                int(document.Comments.Count) == 2,
                f"Word comment count={document.Comments.Count}",
            ),
        ),
        LiveCase(
            "word_live_reply_to_comment",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "comment_index": 1,
                "text": "Verification reply",
                "author": "Capability Verifier",
            },
            lambda _app, document, _payload, _root: (
                int(document.Comments(1).Replies.Count) >= 1,
                f"Word reply count={document.Comments(1).Replies.Count}",
            ),
        ),
        LiveCase(
            "word_live_delete_comment",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "comment_index": 1,
                "delete_replies": True,
            },
            lambda _app, document, _payload, _root: (
                int(document.Comments.Count) == 0,
                f"Word comment count={document.Comments.Count}",
            ),
        ),
        LiveCase(
            "word_live_accept_revisions",
            filename,
            lambda _app, document, _payload, _root: (
                int(document.Revisions.Count) == 0,
                f"Word revision count={document.Revisions.Count}",
            ),
            setup=_add_revision,
        ),
        LiveCase(
            "word_live_reject_revisions",
            filename,
            lambda _app, document, _payload, _root: (
                int(document.Revisions.Count) == 0
                and "Revision marker" not in str(document.Content.Text),
                f"Word revisions={document.Revisions.Count}; marker present={'Revision marker' in str(document.Content.Text)}",
            ),
            setup=_add_revision,
        ),
        LiveCase(
            "word_live_toggle_track_changes",
            lambda _doc, root: {"filename": str(root / "case.docx"), "enable": True},
            lambda _app, document, _payload, _root: (
                bool(document.TrackRevisions),
                f"Word TrackRevisions={bool(document.TrackRevisions)}",
            ),
        ),
        LiveCase(
            "word_live_edit_footnotes_endnotes",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "operation": "add",
                "note_type": "footnote",
                "paragraph_index": 3,
                "text": "Added footnote",
            },
            lambda _app, document, _payload, _root: (
                int(document.Footnotes.Count) == 2,
                f"Word footnote count={document.Footnotes.Count}",
            ),
        ),
        LiveCase(
            "word_live_set_note_configuration",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "note_type": "footnote",
                "starting_number": 1,
                "numbering_rule": "restart_each_section",
                "number_style": "lowercase_roman",
                "location": "beneath_text",
                "separator_text": "Capability separator",
            },
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and int(document.Footnotes.StartingNumber) == 1
                and int(document.Footnotes.NumberingRule) == 1
                and int(document.Footnotes.NumberStyle) == 2
                and int(document.Footnotes.Location) == 1
                and "Capability separator" in str(document.StoryRanges(12).Text),
                (
                    f"Word footnote starting={document.Footnotes.StartingNumber}; "
                    f"rule={document.Footnotes.NumberingRule}; style={document.Footnotes.NumberStyle}; "
                    f"location={document.Footnotes.Location}; "
                    f"separator={document.StoryRanges(12).Text!r}"
                ),
            ),
        ),
        LiveCase(
            "word_live_set_core_properties",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "title": "Verified title",
                "author": "Capability Verifier",
            },
            lambda _app, document, _payload, _root: (
                str(document.BuiltInDocumentProperties("Title").Value) == "Verified title",
                f"Word title={document.BuiltInDocumentProperties('Title').Value!r}",
            ),
        ),
        LiveCase(
            "word_live_save",
            filename,
            lambda _app, document, _payload, _root: (
                bool(document.Saved),
                f"Word Saved={bool(document.Saved)}",
            ),
            setup=lambda _app, document, _root: document.Range(
                document.Content.End - 1, document.Content.End - 1
            ).InsertAfter("Unsaved marker"),
        ),
        LiveCase(
            "word_live_undo",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "times": 1,
                "mcp_only": True,
            },
            _text_absent("Undo target marker"),
            pre_calls=(
                PreCall(
                    "word_live_insert_text",
                    lambda _doc, root: {
                        "filename": str(root / "case.docx"),
                        "text": "Undo target marker",
                        "position": "end",
                    },
                ),
            ),
        ),
        LiveCase(
            "word_live_insert_image",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "image_path": str(image_path),
                "position": "end",
                "width_inches": 1.5,
                "alignment": "center",
            },
            lambda _app, document, _payload, _root: (
                int(document.InlineShapes.Count) + int(document.Shapes.Count) >= 1,
                f"Word inline_shapes={document.InlineShapes.Count}; shapes={document.Shapes.Count}",
            ),
        ),
        LiveCase(
            "word_live_insert_equation",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "equation": "x^2 + y^2 = z^2",
                "position": "end",
                "display_mode": True,
            },
            lambda _app, document, _payload, _root: (
                int(document.OMaths.Count) >= 1,
                f"Word equation count={document.OMaths.Count}",
            ),
        ),
        LiveCase(
            "word_live_set_page_layout",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "section_index": 1,
                "orientation": "landscape",
                "margin_left_inches": 0.8,
            },
            lambda _app, document, _payload, _root: (
                int(document.Sections(1).PageSetup.Orientation) == 1
                and abs(float(document.Sections(1).PageSetup.LeftMargin) - 57.6) < 0.5,
                f"Word orientation={document.Sections(1).PageSetup.Orientation}; left_margin={document.Sections(1).PageSetup.LeftMargin}",
            ),
        ),
        LiveCase(
            "word_live_inspect_layout",
            filename,
            _verify_layout_diagnostics,
            setup=_setup_layout_diagnostics,
        ),
        LiveCase(
            "word_live_add_section_break",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "break_type": "new_page",
                "paragraph_index": 3,
            },
            lambda _app, document, _payload, _root: (
                int(document.Sections.Count) == 2,
                f"Word section count={document.Sections.Count}",
            ),
        ),
        LiveCase(
            "word_live_add_bookmark",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "paragraph_index": 2,
                "bookmark_name": "CapabilityBookmark",
            },
            lambda _app, document, _payload, _root: (
                bool(document.Bookmarks.Exists("CapabilityBookmark")),
                f"Word bookmark exists={bool(document.Bookmarks.Exists('CapabilityBookmark'))}",
            ),
        ),
        LiveCase(
            "word_live_add_watermark",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "text": "VERIFY",
                "section_index": 1,
            },
            lambda _app, document, _payload, _root: (
                int(document.Sections(1).Headers(1).Shapes.Count) >= 1,
                f"Word header shape count={document.Sections(1).Headers(1).Shapes.Count}",
            ),
        ),
        LiveCase(
            "word_live_edit_headers_footers",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "section": 1,
                "story_kind": "footer",
                "variant": "primary",
                "operation": "set",
                "content": "Page {page} of {pages}",
                "alignment": "center",
                "page_number_style": "arabic",
            },
            lambda _app, document, _payload, _root: (
                int(document.Sections(1).Footers(1).Range.Fields.Count) == 2
                and "Page" in str(document.Sections(1).Footers(1).Range.Text),
                f"Word footer fields={document.Sections(1).Footers(1).Range.Fields.Count}; text={document.Sections(1).Footers(1).Range.Text!r}",
            ),
        ),
        LiveCase(
            "word_live_list_fields",
            lambda _doc, root: {"filename": str(root / "case.docx")},
            lambda _app, _document, payload, _root: (
                payload.get("field_count", 0) >= 2
                and any(field.get("type") == "date" for field in payload.get("fields", []))
                and any(
                    field.get("type") == "index_entry" and field.get("result_available") is False
                    for field in payload.get("fields", [])
                ),
                f"listed fields={payload.get('field_count')!r}",
            ),
            setup=_add_fields_for_inspection,
        ),
        LiveCase(
            "word_live_update_fields",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "field_indices": [1],
            },
            lambda _app, document, payload, _root: (
                payload.get("updated_count") == 1
                and bool(str(document.Fields(1).Result.Text).strip()),
                (
                    f"updated fields={payload.get('updated_count')!r}; "
                    f"result={str(document.Fields(1).Result.Text)!r}"
                ),
            ),
            setup=_add_date_field,
        ),
        LiveCase(
            "word_live_unlink_fields",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "field_indices": [1, 2],
            },
            lambda _app, document, payload, _root: (
                payload.get("unlinked_count") == 2
                and int(document.Fields.Count) == 0
                and bool(str(document.Content.Text).strip()),
                (
                    f"unlinked fields={payload.get('unlinked_count')!r}; "
                    f"Word main fields={document.Fields.Count}"
                ),
            ),
            setup=_add_unlinkable_fields,
        ),
        LiveCase(
            "word_live_list_tables_of_contents",
            lambda _doc, root: {"filename": str(root / "case.docx")},
            lambda _app, document, payload, _root: (
                payload.get("toc_count") == 1 and int(document.TablesOfContents.Count) == 1,
                f"payload TOCs={payload.get('toc_count')!r}; Word TOCs={document.TablesOfContents.Count}",
            ),
            setup=_add_table_of_contents,
        ),
        LiveCase(
            "word_live_create_table_of_contents",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "target": "document_start",
                "upper_heading_level": 1,
                "lower_heading_level": 3,
            },
            lambda _app, document, payload, _root: (
                payload.get("success") is True and int(document.TablesOfContents.Count) == 1,
                f"Word TOCs={document.TablesOfContents.Count}",
            ),
        ),
        LiveCase(
            "word_live_update_table_of_contents",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "toc_index": 1,
                "mode": "all",
            },
            lambda _app, document, payload, _root: (
                payload.get("success") is True and int(document.TablesOfContents.Count) == 1,
                f"Word TOCs after update={document.TablesOfContents.Count}",
            ),
            setup=_add_table_of_contents,
        ),
        LiveCase(
            "word_live_delete_table_of_contents",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "toc_index": 1,
            },
            lambda _app, document, payload, _root: (
                payload.get("remaining_toc_count") == 0
                and int(document.TablesOfContents.Count) == 0,
                f"Word TOCs after delete={document.TablesOfContents.Count}",
            ),
            setup=_add_table_of_contents,
        ),
        LiveCase(
            "word_live_list_hyperlinks",
            lambda _doc, root: {"filename": str(root / "case.docx")},
            lambda _app, document, payload, _root: (
                payload.get("hyperlink_count") == 1 and int(document.Hyperlinks.Count) == 1,
                f"payload hyperlinks={payload.get('hyperlink_count')!r}",
            ),
            setup=_add_hyperlink,
        ),
        LiveCase(
            "word_live_add_hyperlink",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "address": "https://example.com/capability",
                "display_text": "Capability link",
            },
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and int(document.Hyperlinks.Count) == 1
                and "Capability link" in str(document.Content.Text),
                f"Word hyperlinks={document.Hyperlinks.Count}",
            ),
        ),
        LiveCase(
            "word_live_update_hyperlink",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "hyperlink_index": 1,
                "address": "https://example.com/updated",
            },
            lambda _app, document, payload, _root: (
                payload.get("success") is True
                and "example.com/updated" in str(document.Hyperlinks(1).Address),
                f"Word address={document.Hyperlinks(1).Address!r}",
            ),
            setup=_add_hyperlink,
        ),
        LiveCase(
            "word_live_remove_hyperlink",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "hyperlink_index": 1,
            },
            lambda _app, document, payload, _root: (
                payload.get("remaining_hyperlinks") == 0
                and int(document.Hyperlinks.Count) == 0
                and "Beta marker" in str(document.Content.Text),
                f"Word hyperlinks={document.Hyperlinks.Count}; text preserved={'Beta marker' in str(document.Content.Text)}",
            ),
            setup=_add_hyperlink,
        ),
        LiveCase(
            "word_live_list_cross_reference_targets",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "reference_type": "heading",
            },
            lambda _app, _document, payload, _root: (
                payload.get("count", 0) >= 1
                and any(
                    "Alpha marker" in target.get("label", "")
                    for target in payload.get("targets", [])
                ),
                f"native heading targets={payload.get('count')!r}",
            ),
        ),
        LiveCase(
            "word_live_insert_cross_reference",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "reference_type": "heading",
                "target_index": 1,
                "reference_kind": "content_text",
                "position": "end",
            },
            lambda _app, document, payload, _root: (
                payload.get("success") is True and int(document.Fields.Count) >= 1,
                f"Word main-story fields={document.Fields.Count}",
            ),
        ),
        LiveCase(
            "word_live_capture_pages",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "pages": [1],
                "dpi": 96,
            },
            lambda _app, _document, payload, _root: (
                [item.get("page") for item in payload.get("pages", [])] == [1],
                f"captured_pages={[item.get('page') for item in payload.get('pages', [])]!r}",
            ),
        ),
        LiveCase(
            "word_live_edit_batch",
            lambda _doc, root: {
                "filename": str(root / "case.docx"),
                "operations": [
                    {
                        "tool": "word_live_replace_text",
                        "arguments": {
                            "find_text": "Alpha marker",
                            "replace_text": "Batch verified",
                        },
                    }
                ],
                "contains_text": ["Batch verified"],
                "absent_text": ["Alpha marker"],
                "capture_pages": [1],
            },
            _text_contains("Batch verified"),
        ),
        LiveCase(
            "word_live_insert_file",
            lambda _doc, root: {
                "source_path": str(root / "merge-source.docx"),
                "filename": str(root / "case.docx"),
                "position": "end",
            },
            _verify_native_merge,
            setup=_make_merge_source,
        ),
        LiveCase(
            "word_live_create_document_snapshot",
            lambda _doc, root: {
                "snapshot_path": str(root / "snapshot.json"),
                "filename": str(root / "case.docx"),
            },
            _verify_snapshot_capture,
            setup=_setup_saved_document_with_linked_header,
        ),
        LiveCase(
            "word_live_diff_document_snapshots",
            lambda _doc, root: {
                "before_snapshot_path": str(root / "before.json"),
                "after_snapshot_path": str(root / "after.json"),
            },
            lambda _app, _document, payload, _root: (
                payload.get("success") is True
                and payload.get("same_source_path") is True
                and payload.get("identical_semantic_content") is False
                and (
                    payload.get("paragraph_change_group_count", 0)
                    + payload.get("component_leaf_change_count", 0)
                    >= 1
                )
                and len(payload.get("paragraph_operations", [])) >= 1,
                "paragraph_change_groups="
                f"{payload.get('paragraph_change_group_count')!r}; "
                "component_leaf_changes="
                f"{payload.get('component_leaf_change_count')!r}; "
                f"paragraph_operations={len(payload.get('paragraph_operations', []))}",
            ),
            pre_calls=(
                PreCall(
                    "word_live_create_document_snapshot",
                    lambda _doc, root: {
                        "snapshot_path": str(root / "before.json"),
                        "filename": str(root / "case.docx"),
                    },
                ),
                PreCall(
                    "word_live_replace_text",
                    lambda _doc, root: {
                        "filename": str(root / "case.docx"),
                        "find_text": "Beta marker",
                        "replace_text": "Snapshot changed marker",
                    },
                ),
                PreCall(
                    "word_live_create_document_snapshot",
                    lambda _doc, root: {
                        "snapshot_path": str(root / "after.json"),
                        "filename": str(root / "case.docx"),
                    },
                ),
            ),
        ),
        LiveCase(
            "word_live_navigate",
            lambda _doc, root: {"filename": str(root / "case.docx"), "page": 2},
            _verify_navigation,
            setup=_setup_navigation,
        ),
    ]


async def _run_file_cases(client: Client, root: Path) -> list[ResultRecord]:
    records: list[ResultRecord] = []
    source = root / "closed-source.docx"
    _make_closed_docx(source)
    cases: list[tuple[str, dict[str, Any], Callable[[dict[str, Any]], tuple[bool, str]]]] = [
        (
            "create_document",
            {"filename": str(root / "created.docx"), "title": "Created verification"},
            lambda payload: (
                payload.get("success") is True and (root / "created.docx").is_file(),
                f"created file exists={(root / 'created.docx').is_file()}",
            ),
        ),
        (
            "copy_document",
            {
                "source_filename": str(source),
                "destination_filename": str(root / "copied.docx"),
            },
            lambda payload: (
                payload.get("success") is True and (root / "copied.docx").is_file(),
                f"copied file exists={(root / 'copied.docx').is_file()}",
            ),
        ),
        (
            "get_document_info",
            {"filename": str(source)},
            lambda payload: (
                payload.get("document") == str(source.resolve())
                and payload.get("body_paragraph_count", 0) >= 2,
                (
                    f"document={payload.get('document')!r}; "
                    f"body_paragraphs={payload.get('body_paragraph_count')!r}"
                ),
            ),
        ),
        (
            "list_available_documents",
            {"directory": str(root), "recursive": False},
            lambda payload: (
                payload.get("success") is True and payload.get("count", 0) >= 3,
                f"listed count={payload.get('count')!r}",
            ),
        ),
        (
            "convert_to_pdf",
            {"filename": str(source), "output_filename": str(root / "converted.pdf")},
            lambda _payload: (
                (root / "converted.pdf").is_file() and (root / "converted.pdf").stat().st_size > 0,
                (
                    f"PDF exists={(root / 'converted.pdf').is_file()}; "
                    f"bytes={(root / 'converted.pdf').stat().st_size if (root / 'converted.pdf').is_file() else 0}"
                ),
            ),
        ),
    ]

    for name, arguments, verify in cases:
        started = time.perf_counter()
        try:
            result, payload, response_text, nested_error = await _invoke(client, name, arguments)
            post_ok, evidence = verify(payload)
            mcp_error = bool(result.is_error)
            status = "PASS" if not mcp_error and nested_error is None and post_ok else "FAIL"
            records.append(
                ResultRecord(
                    tool=name,
                    status=status,
                    invocation="FastMCP Client (in-memory transport)",
                    elapsed_seconds=round(time.perf_counter() - started, 3),
                    mcp_is_error=mcp_error,
                    nested_error=nested_error,
                    postcondition=evidence,
                    response_excerpt=response_text[:1000],
                    arguments={key: _json_safe(value) for key, value in arguments.items()},
                )
            )
        except Exception as exc:
            records.append(
                ResultRecord(
                    tool=name,
                    status="FAIL",
                    invocation="FastMCP Client (in-memory transport)",
                    elapsed_seconds=round(time.perf_counter() - started, 3),
                    mcp_is_error=False,
                    nested_error=f"{type(exc).__name__}: {exc}",
                    postcondition="not evaluated",
                    response_excerpt="",
                    arguments={key: _json_safe(value) for key, value in arguments.items()},
                )
            )
    return records


async def _run_open_case(client: Client, app: Any, root: Path) -> ResultRecord:
    path = root / "open-target.docx"
    document = _make_live_document(app, path)
    document.Close(SaveChanges=False)
    arguments = {"filename": str(path)}
    started = time.perf_counter()
    opened = None
    try:
        result, payload, response_text, nested_error = await _invoke(
            client, "word_live_open", arguments
        )
        opened = word_com.find_document(app, str(path))
        post_ok = opened is not None and str(opened.FullName).casefold() == str(path).casefold()
        status = "PASS" if not result.is_error and nested_error is None and post_ok else "FAIL"
        return ResultRecord(
            tool="word_live_open",
            status=status,
            invocation="FastMCP Client (in-memory transport)",
            elapsed_seconds=round(time.perf_counter() - started, 3),
            mcp_is_error=bool(result.is_error),
            nested_error=nested_error,
            postcondition=f"Word opened exact path={post_ok}",
            response_excerpt=response_text[:1000],
            arguments=arguments,
        )
    except Exception as exc:
        return ResultRecord(
            tool="word_live_open",
            status="FAIL",
            invocation="FastMCP Client (in-memory transport)",
            elapsed_seconds=round(time.perf_counter() - started, 3),
            mcp_is_error=False,
            nested_error=f"{type(exc).__name__}: {exc}",
            postcondition="not evaluated",
            response_excerpt="",
            arguments=arguments,
        )
    finally:
        if opened is not None:
            opened.Close(SaveChanges=False)


async def _run_close_case(client: Client, app: Any, root: Path) -> ResultRecord:
    path = root / "close-target.docx"
    document = _make_live_document(app, path)
    arguments = {"filename": str(path)}
    started = time.perf_counter()
    try:
        result, payload, response_text, nested_error = await _invoke(
            client, "word_live_close", arguments
        )
        still_open = any(
            str(app.Documents(index).FullName).casefold() == str(path).casefold()
            for index in range(1, int(app.Documents.Count) + 1)
        )
        post_ok = (
            payload.get("success") is True
            and not still_open
            and path.is_file()
            and payload.get("save_mode") == "require_saved"
        )
        return ResultRecord(
            tool="word_live_close",
            status=("PASS" if not result.is_error and nested_error is None and post_ok else "FAIL"),
            invocation="FastMCP Client (in-memory transport) + isolated Word COM",
            elapsed_seconds=round(time.perf_counter() - started, 3),
            mcp_is_error=bool(result.is_error),
            nested_error=nested_error,
            postcondition=f"closed={not still_open}; saved file exists={path.is_file()}",
            response_excerpt=response_text[:1000],
            arguments=arguments,
        )
    except Exception as exc:
        try:
            document.Close(SaveChanges=False)
        except Exception as exc:
            logger.warning("Could not close document after close-tool case: %s", exc)
        return ResultRecord(
            tool="word_live_close",
            status="FAIL",
            invocation="FastMCP Client (in-memory transport) + isolated Word COM",
            elapsed_seconds=round(time.perf_counter() - started, 3),
            mcp_is_error=False,
            nested_error=f"{type(exc).__name__}: {exc}",
            postcondition="not evaluated",
            response_excerpt="",
            arguments=arguments,
        )


async def _run_rename_case(client: Client, app: Any, root: Path) -> ResultRecord:
    original = root / "rename-source.docx"
    destination = root / "rename-destination.docx"
    document = _make_live_document(app, original)
    arguments = {"filename": str(original), "new_path": str(destination)}
    started = time.perf_counter()
    try:
        result, payload, response_text, nested_error = await _invoke(
            client, "word_live_rename", arguments
        )
        actual_path = Path(str(document.FullName)).resolve()
        post_ok = (
            payload.get("success") is True
            and actual_path == destination.resolve()
            and destination.is_file()
            and not original.exists()
            and "Alpha marker" in str(document.Content.Text)
            and bool(document.Saved)
        )
        return ResultRecord(
            tool="word_live_rename",
            status=("PASS" if not result.is_error and nested_error is None and post_ok else "FAIL"),
            invocation="FastMCP Client (in-memory transport) + isolated Word COM",
            elapsed_seconds=round(time.perf_counter() - started, 3),
            mcp_is_error=bool(result.is_error),
            nested_error=nested_error,
            postcondition=(
                f"active path={actual_path}; destination exists={destination.is_file()}; "
                f"source removed={not original.exists()}; content preserved="
                f"{'Alpha marker' in str(document.Content.Text)}"
            ),
            response_excerpt=response_text[:1000],
            arguments=arguments,
        )
    except Exception as exc:
        return ResultRecord(
            tool="word_live_rename",
            status="FAIL",
            invocation="FastMCP Client (in-memory transport) + isolated Word COM",
            elapsed_seconds=round(time.perf_counter() - started, 3),
            mcp_is_error=False,
            nested_error=f"{type(exc).__name__}: {exc}",
            postcondition="not evaluated",
            response_excerpt="",
            arguments=arguments,
        )
    finally:
        try:
            document.Close(SaveChanges=False)
        except Exception as exc:
            logger.warning("Could not close document after rename-tool case: %s", exc)


async def _run_live_cases(client: Client, app: Any, root: Path, image: Path) -> list[ResultRecord]:
    records: list[ResultRecord] = []
    for case in _live_cases(image):
        case_root = root / case.name
        case_root.mkdir()
        path = case_root / "case.docx"
        document = None
        started = time.perf_counter()
        arguments: dict[str, Any] = {}
        try:
            document = _make_live_document(app, path)
            if case.setup is not None:
                case.setup(app, document, case_root)

            for pre_call in case.pre_calls:
                pre_arguments = pre_call.arguments(document, case_root)
                pre_result, _pre_payload, _pre_text, pre_nested_error = await _invoke(
                    client, pre_call.name, pre_arguments
                )
                if pre_result.is_error or pre_nested_error is not None:
                    raise RuntimeError(
                        f"pre-call {pre_call.name} failed: {pre_nested_error or _pre_text}"
                    )

            arguments = case.arguments(document, case_root)
            result, payload, response_text, nested_error = await _invoke(
                client, case.name, arguments
            )
            post_ok, evidence = case.verifier(app, document, payload, case_root)
            mcp_error = bool(result.is_error)
            status = "PASS" if not mcp_error and nested_error is None and post_ok else "FAIL"
            image_count = sum(
                1 for item in result.content if getattr(item, "type", None) == "image"
            )
            if case.name in {"word_live_capture_pages", "word_live_edit_batch"}:
                post_ok = post_ok and image_count >= 1
                evidence = f"{evidence}; MCP images={image_count}"
                status = "PASS" if not mcp_error and nested_error is None and post_ok else "FAIL"
            records.append(
                ResultRecord(
                    tool=case.name,
                    status=status,
                    invocation="FastMCP Client (in-memory transport) + isolated Word COM",
                    elapsed_seconds=round(time.perf_counter() - started, 3),
                    mcp_is_error=mcp_error,
                    nested_error=nested_error,
                    postcondition=evidence,
                    response_excerpt=response_text[:1000],
                    arguments={key: _json_safe(value) for key, value in arguments.items()},
                )
            )
        except Exception as exc:
            records.append(
                ResultRecord(
                    tool=case.name,
                    status="FAIL",
                    invocation="FastMCP Client (in-memory transport) + isolated Word COM",
                    elapsed_seconds=round(time.perf_counter() - started, 3),
                    mcp_is_error=False,
                    nested_error=f"{type(exc).__name__}: {exc}",
                    postcondition="not evaluated",
                    response_excerpt="",
                    arguments={key: _json_safe(value) for key, value in arguments.items()},
                )
            )
        finally:
            if document is not None:
                try:
                    document.Close(SaveChanges=False)
                except Exception as exc:
                    logger.warning("Could not close live-case document %s: %s", case.name, exc)
    return records


async def _require_call(
    client: Client, name: str, arguments: dict[str, Any]
) -> tuple[Any, dict[str, Any]]:
    result, payload, response_text, nested_error = await _invoke(client, name, arguments)
    if result.is_error or nested_error is not None:
        raise RuntimeError(f"{name} failed: {nested_error or response_text}")
    return result, payload


async def _run_stateful_document_workflow(client: Client, app: Any, root: Path) -> WorkflowRecord:
    """Exercise cumulative edits, implicit targeting, save, close, and reopen."""
    workflow = "stateful single-document edit, save, and reopen"
    workflow_root = root / "workflow-stateful"
    workflow_root.mkdir()
    path = workflow_root / "stateful.docx"
    document = None
    steps = 0
    expected_steps = 14
    started = time.perf_counter()
    try:
        document = _make_live_document(app, path)
        document.Activate()

        _result, payload = await _require_call(client, "word_live_get_text", {})
        if not any(
            "Alpha marker" in item.get("text", "") for item in payload.get("paragraphs", [])
        ):
            raise RuntimeError("implicit active-document read did not return Alpha marker")
        steps += 1

        await _require_call(
            client,
            "word_live_edit_batch",
            {
                "operations": [
                    {
                        "tool": "word_live_replace_text",
                        "arguments": {
                            "find_text": "Beta marker",
                            "replace_text": "Workflow beta",
                        },
                    },
                    {
                        "tool": "word_live_insert_text",
                        "arguments": {"text": "Workflow inserted", "position": "end"},
                    },
                ],
                "contains_text": ["Workflow beta", "Workflow inserted"],
                "absent_text": ["Beta marker"],
                "capture_pages": [1],
            },
        )
        steps += 1

        calls = [
            (
                "word_live_format_text",
                {"start_paragraph": 1, "end_paragraph": 1, "bold": True, "font_size": 16},
            ),
            (
                "word_live_add_comment",
                {
                    "paragraph_index": 3,
                    "text": "Workflow comment",
                    "author": "Capability Verifier",
                },
            ),
            (
                "word_live_edit_footnotes_endnotes",
                {
                    "operation": "add",
                    "note_type": "endnote",
                    "paragraph_index": 2,
                    "text": "Workflow endnote",
                },
            ),
            (
                "word_live_edit_headers_footers",
                {
                    "section": 1,
                    "story_kind": "footer",
                    "variant": "primary",
                    "operation": "set",
                    "content": "Workflow {page} / {pages}",
                    "alignment": "center",
                },
            ),
            (
                "word_live_add_bookmark",
                {"paragraph_index": 2, "bookmark_name": "WorkflowBookmark"},
            ),
            (
                "word_live_set_core_properties",
                {"title": "Workflow verified", "author": "Capability Verifier"},
            ),
            (
                "word_live_add_table",
                {
                    "rows": 2,
                    "cols": 2,
                    "data": [["Flow A1", "Flow A2"], ["Flow B1", "Flow B2"]],
                    "position": "end",
                },
            ),
        ]
        for name, arguments in calls:
            await _require_call(client, name, arguments)
            steps += 1

        _save_result, save_payload = await _require_call(client, "word_live_save", {})
        if Path(str(save_payload.get("path", ""))).resolve() != path.resolve():
            raise RuntimeError(f"save targeted the wrong document: {save_payload}")
        saved_flag_after_call = bool(document.Saved)
        steps += 1

        preclose_checks = {
            "text": "Workflow beta" in str(document.Content.Text)
            and "Workflow inserted" in str(document.Content.Text),
            "format": bool(document.Paragraphs(1).Range.Font.Bold)
            and abs(float(document.Paragraphs(1).Range.Font.Size) - 16) < 0.1,
            "comments": int(document.Comments.Count) == 2,
            "endnotes": int(document.Endnotes.Count) == 1,
            "footer_fields": int(document.Sections(1).Footers(1).Range.Fields.Count) == 2,
            "bookmark": bool(document.Bookmarks.Exists("WorkflowBookmark")),
            "title": str(document.BuiltInDocumentProperties("Title").Value) == "Workflow verified",
            "tables": int(document.Tables.Count) == 2,
        }
        if not all(preclose_checks.values()):
            raise RuntimeError(f"pre-close accumulated state failed: {preclose_checks}")
        steps += 1

        # Close through the public lifecycle tool with an explicit save policy.
        # Word can dirty fields asynchronously after an earlier save, so a complete
        # edit workflow should save at the lifecycle boundary as well.
        await _require_call(
            client,
            "word_live_close",
            {"filename": str(path), "save_mode": "save"},
        )
        document = None
        await _require_call(client, "word_live_open", {"filename": str(path)})
        document = word_com.find_document(app, str(path))
        steps += 1

        await _require_call(client, "word_live_get_info", {})
        steps += 1

        persisted_checks = {
            "text": "Workflow beta" in str(document.Content.Text)
            and "Workflow inserted" in str(document.Content.Text),
            "comments": int(document.Comments.Count) == 2,
            "endnotes": int(document.Endnotes.Count) == 1,
            "footer_fields": int(document.Sections(1).Footers(1).Range.Fields.Count) == 2,
            "bookmark": bool(document.Bookmarks.Exists("WorkflowBookmark")),
            "title": str(document.BuiltInDocumentProperties("Title").Value) == "Workflow verified",
            "tables": int(document.Tables.Count) == 2,
        }
        if not all(persisted_checks.values()):
            raise RuntimeError(f"reopen persistence failed: {persisted_checks}")
        steps += 1

        return WorkflowRecord(
            workflow=workflow,
            status="PASS",
            elapsed_seconds=round(time.perf_counter() - started, 3),
            steps_completed=steps,
            expected_steps=expected_steps,
            evidence=(
                "All cumulative text, formatting, comment, endnote, footer field, bookmark, "
                "property, table, save, public close-with-save, reopen, and implicit "
                f"active-document checks passed; immediate Word Saved flag={saved_flag_after_call}"
            ),
        )
    except Exception as exc:
        return WorkflowRecord(
            workflow=workflow,
            status="FAIL",
            elapsed_seconds=round(time.perf_counter() - started, 3),
            steps_completed=steps,
            expected_steps=expected_steps,
            evidence="Workflow stopped before all postconditions passed",
            error=f"{type(exc).__name__}: {exc}",
        )
    finally:
        if document is not None:
            try:
                document.Close(SaveChanges=False)
            except Exception as exc:
                logger.warning("Could not close stateful-workflow document: %s", exc)


def _make_simple_live_document(app: Any, path: Path, marker: str) -> Any:
    document = app.Documents.Add()
    document.Content.Text = f"{marker}\r"
    document.SaveAs2(str(path), AddToRecentFiles=False)
    return document


async def _run_multiple_document_workflow(client: Client, app: Any, root: Path) -> WorkflowRecord:
    """Verify active, named, and full-path selection with several open documents."""
    workflow = "multiple open documents and target selection"
    workflow_root = root / "workflow-multiple"
    workflow_root.mkdir()
    paths = [workflow_root / f"document-{index}.docx" for index in range(1, 4)]
    documents: list[Any] = []
    steps = 0
    expected_steps = 6
    started = time.perf_counter()
    try:
        for index, path in enumerate(paths, 1):
            documents.append(_make_simple_live_document(app, path, f"Document {index} marker"))
        documents[2].Activate()

        _result, payload = await _require_call(client, "word_live_get_text", {})
        if "Document 3 marker" not in json.dumps(payload):
            raise RuntimeError("implicit target was not the active third document")
        steps += 1

        _result, payload = await _require_call(
            client, "word_live_get_text", {"filename": paths[0].name}
        )
        if "Document 1 marker" not in json.dumps(payload):
            raise RuntimeError("unique basename did not select the first document")
        steps += 1

        _result, payload = await _require_call(
            client, "word_live_get_text", {"filename": str(paths[1])}
        )
        if "Document 2 marker" not in json.dumps(payload):
            raise RuntimeError("full path did not select the second document")
        steps += 1

        await _require_call(
            client,
            "word_live_insert_text",
            {"filename": str(paths[0]), "text": "First only", "position": "end"},
        )
        if "First only" not in str(documents[0].Content.Text):
            raise RuntimeError("explicit edit did not reach the first document")
        if any("First only" in str(document.Content.Text) for document in documents[1:]):
            raise RuntimeError("explicit edit leaked into another open document")
        steps += 1

        _result, payload = await _require_call(client, "word_live_list_open", {})
        if payload.get("count", 0) < 3:
            raise RuntimeError(f"list_open reported fewer than three documents: {payload}")
        steps += 1

        documents[1].Activate()
        _result, payload = await _require_call(client, "word_live_get_text", {})
        if "Document 2 marker" not in json.dumps(payload):
            raise RuntimeError("active target did not follow Word activation change")
        steps += 1

        return WorkflowRecord(
            workflow=workflow,
            status="PASS",
            elapsed_seconds=round(time.perf_counter() - started, 3),
            steps_completed=steps,
            expected_steps=expected_steps,
            evidence=(
                "Active-document, unique-basename, full-path, isolated mutation, listing, "
                "and activation-switch checks passed across three simultaneously open documents"
            ),
        )
    except Exception as exc:
        return WorkflowRecord(
            workflow=workflow,
            status="FAIL",
            elapsed_seconds=round(time.perf_counter() - started, 3),
            steps_completed=steps,
            expected_steps=expected_steps,
            evidence="Workflow stopped before all selection checks passed",
            error=f"{type(exc).__name__}: {exc}",
        )
    finally:
        for document in reversed(documents):
            try:
                document.Close(SaveChanges=False)
            except Exception as exc:
                logger.warning("Could not close multiple-document workflow document: %s", exc)


def _write_reports(
    records: list[ResultRecord],
    workflows: list[WorkflowRecord],
    metadata: dict[str, Any],
    output: Path,
) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    json_output = output.with_suffix(".json")
    document = {
        "metadata": metadata,
        "summary": {
            "total": len(records),
            "passed": sum(record.status == "PASS" for record in records),
            "failed": sum(record.status == "FAIL" for record in records),
            "not_run": sum(record.status == "NOT_RUN" for record in records),
            "workflow_total": len(workflows),
            "workflow_passed": sum(workflow.status == "PASS" for workflow in workflows),
            "workflow_failed": sum(workflow.status == "FAIL" for workflow in workflows),
        },
        "results": [asdict(record) for record in records],
        "workflows": [asdict(workflow) for workflow in workflows],
    }
    json_output.write_text(json.dumps(document, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# Word integration results",
        "",
        f"Generated: {metadata['generated_at']}",
        "",
        "This report records actual calls through `fastmcp.Client` using its in-memory",
        "transport. Live cases used a private, hidden Microsoft Word COM instance and",
        "disposable documents. It does not claim stdio/HTTP transport verification.",
        "Each tool receives one explicit representative case; PASS does not mean every",
        "argument combination or multi-operation branch has been verified.",
        "",
        f"- Registered tools: {metadata['registered_tools']}",
        f"- Cases run: {len(records)}",
        f"- Passed: {document['summary']['passed']}",
        f"- Failed: {document['summary']['failed']}",
        f"- Not run: {document['summary']['not_run']}",
        f"- Stateful workflows passed: {document['summary']['workflow_passed']} / {document['summary']['workflow_total']}",
        f"- Word version: {metadata.get('word_version')}",
        f"- Python: {metadata['python']}",
        f"- Package version: {metadata.get('package_version')}",
        f"- Git commit: {metadata.get('git_commit')}",
        f"- Git dirty: {metadata.get('git_dirty')}",
        f"- Source fingerprint: `{metadata.get('source_fingerprint_sha256')}`",
        f"- Integration runner SHA-256: `{metadata.get('runner_sha256')}`",
        "",
        "| Tool | Status | MCP error | Nested error | Postcondition |",
        "| --- | --- | --- | --- | --- |",
    ]
    for record in records:
        nested = (record.nested_error or "").replace("|", "\\|").replace("\n", " ")
        postcondition = record.postcondition.replace("|", "\\|").replace("\n", " ")
        lines.append(
            f"| `{record.tool}` | {record.status} | {record.mcp_is_error} | "
            f"{nested} | {postcondition} |"
        )
    lines.extend(
        [
            "",
            "## Stateful workflow results",
            "",
            "| Workflow | Status | Steps | Evidence | Error |",
            "| --- | --- | --- | --- | --- |",
        ]
    )
    for workflow in workflows:
        evidence = workflow.evidence.replace("|", "\\|").replace("\n", " ")
        error = (workflow.error or "").replace("|", "\\|").replace("\n", " ")
        lines.append(
            f"| {workflow.workflow} | {workflow.status} | "
            f"{workflow.steps_completed}/{workflow.expected_steps} | {evidence} | {error} |"
        )
    lines.extend(
        [
            "",
            "Full arguments, response excerpts, timings, and errors are in "
            f"[`{json_output.name}`]({json_output.name}).",
            "",
        ]
    )
    output.write_text("\n".join(lines), encoding="utf-8")


async def run(output: Path) -> int:
    server = create_server(tool_mode="full")
    registered_names = [tool.name for tool in await server.list_tools()]
    records: list[ResultRecord] = []
    workflows: list[WorkflowRecord] = []
    app = None
    word_version = "unknown"
    with tempfile.TemporaryDirectory(prefix="word-mcp-integration-") as temp:
        root = Path(temp)
        image = root / "verification.png"
        _make_png(image)
        try:
            app = win32com.client.DispatchEx("Word.Application")
            app.Visible = False
            app.DisplayAlerts = 0
            word_version = str(app.Version)
            word_com.remember_word_app(app)

            async with Client(server) as client:
                listed = await client.list_tools()
                listed_names = {tool.name for tool in listed}
                if listed_names != set(registered_names):
                    raise RuntimeError(
                        "MCP client tool list differs from provider discovery: "
                        f"missing={sorted(set(registered_names) - listed_names)}, "
                        f"extra={sorted(listed_names - set(registered_names))}"
                    )
                records.extend(await _run_file_cases(client, root))
                records.append(await _run_open_case(client, app, root))
                records.append(await _run_close_case(client, app, root))
                records.append(await _run_rename_case(client, app, root))
                records.extend(await _run_live_cases(client, app, root, image))
                workflows.append(await _run_stateful_document_workflow(client, app, root))
                workflows.append(await _run_multiple_document_workflow(client, app, root))

            tested = {record.tool for record in records}
            for name in registered_names:
                if name not in tested:
                    records.append(
                        ResultRecord(
                            tool=name,
                            status="NOT_RUN",
                            invocation="none",
                            elapsed_seconds=0,
                            mcp_is_error=False,
                            nested_error="No verification case defined",
                            postcondition="not evaluated",
                            response_excerpt="",
                        )
                    )
        finally:
            if app is not None:
                try:
                    while app.Documents.Count:
                        app.Documents(1).Close(SaveChanges=False)
                except Exception as exc:
                    logger.warning("Could not close all integration documents: %s", exc)
                try:
                    app.Quit(SaveChanges=False)
                except Exception as exc:
                    logger.warning("Could not quit integration Word instance: %s", exc)
            word_com._WORD_APP = None

    records.sort(key=lambda item: registered_names.index(item.tool))
    repository = Path(__file__).resolve().parents[2]
    fingerprint_paths = [
        *sorted((repository / "src").rglob("*.py")),
        *sorted((repository / "tests").rglob("*.py")),
        *sorted((repository / ".github" / "ci").rglob("*.py")),
        repository / "pyproject.toml",
        repository / "uv.lock",
    ]
    fingerprint = hashlib.sha256()
    for path in fingerprint_paths:
        fingerprint.update(path.relative_to(repository).as_posix().encode())
        fingerprint.update(path.read_bytes())
    try:
        git_commit = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            cwd=repository,
            check=True,
            capture_output=True,
            text=True,
        ).stdout.strip()
        git_dirty = bool(
            subprocess.run(
                ["git", "status", "--porcelain"],
                cwd=repository,
                check=True,
                capture_output=True,
                text=True,
            ).stdout.strip()
        )
    except Exception:
        git_commit = "unavailable"
        git_dirty = None
    metadata = {
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S %z"),
        "registered_tools": len(registered_names),
        "python": sys.version.replace("\n", " "),
        "platform": platform.platform(),
        "word_version": word_version,
        "package_version": version("word-mcp-codemode-live"),
        "fastmcp_version": version("fastmcp"),
        "mcp_version": version("mcp"),
        "git_commit": git_commit,
        "git_dirty": git_dirty,
        "source_fingerprint_sha256": fingerprint.hexdigest(),
        "runner_sha256": hashlib.sha256(Path(__file__).read_bytes()).hexdigest(),
        "transport": "FastMCP Client in-memory transport",
        "temporary_documents_retained": False,
    }
    _write_reports(records, workflows, metadata, output)
    failed = sum(record.status == "FAIL" for record in records)
    not_run = sum(record.status == "NOT_RUN" for record in records)
    workflow_failed = sum(workflow.status == "FAIL" for workflow in workflows)
    print(
        f"Word integration: {len(records)} total, {len(records) - failed - not_run} passed, {failed} failed, {not_run} not run"
    )
    print(f"Markdown report: {output}")
    print(f"JSON report: {output.with_suffix('.json')}")
    print(
        f"Stateful workflows: {len(workflows) - workflow_failed} passed, {workflow_failed} failed"
    )
    return 1 if failed or not_run or workflow_failed else 0


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output",
        type=Path,
        default=Path(tempfile.gettempdir()) / "word-mcp-live-integration.md",
        help="Markdown report path; a JSON report is written beside it.",
    )
    args = parser.parse_args()
    raise SystemExit(asyncio.run(run(args.output.resolve())))


if __name__ == "__main__":
    main()
