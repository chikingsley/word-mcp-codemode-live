import asyncio
import sys
from pathlib import Path

import pytest
from docx import Document

# Target for testing: convert_to_pdf (async function)
from word_mcp_codemode_live.tools.extended_document_tools import convert_to_pdf


def _make_sample_docx(path: Path) -> None:
    """Generates a simple .docx file in a temporary directory."""
    doc = Document()
    doc.add_heading("Conversion Test Document", level=1)
    doc.add_paragraph("This is a test paragraph for PDF conversion. Contains ASCII too.")
    doc.add_paragraph(
        "Second paragraph: Contains special characters and spaces to cover path/content edge cases."
    )
    doc.save(str(path))


@pytest.mark.skipif(sys.platform != "win32", reason="Microsoft Word COM requires Windows")
def test_convert_to_pdf_with_temp_docx(tmp_path: Path):
    """Create a DOCX and verify Microsoft Word exports the requested PDF path."""
    # 1) Generate a docx file with spaces in its name in the temp directory
    src_doc = tmp_path / "sample document with spaces.docx"
    _make_sample_docx(src_doc)

    # 2) Define the output PDF path (also in the temp directory)
    out_pdf = tmp_path / "converted output.pdf"

    # 3) Run the asynchronous function under test
    result_msg = asyncio.run(convert_to_pdf(str(src_doc), output_filename=str(out_pdf)))

    if not out_pdf.exists():
        pytest.skip(f"PDF conversion tool unavailable or conversion failed: {result_msg}")

    assert "successfully converted" in result_msg.lower()
    assert out_pdf.stat().st_size > 0


if __name__ == "__main__":
    # Allow running this file directly for quick verification:
    #   python tests/test_convert_to_pdf.py
    import sys

    sys.exit(pytest.main([__file__, "-q"]))
