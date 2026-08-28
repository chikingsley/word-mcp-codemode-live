from pathlib import Path

import pytest
from docx import Document

from word_mcp_codemode_live.tools.content_tools import add_heading, add_paragraph
from word_mcp_codemode_live.tools.document_tools import create_document


@pytest.mark.asyncio
async def test_paragraph_and_heading_formatting(tmp_path: Path) -> None:
    document_path = tmp_path / "formatting.docx"

    await create_document(str(document_path), title="Formatting Test", author="Test Suite")
    await add_paragraph(
        str(document_path),
        "JAMES MEHORTER",
        font_name="Helvetica",
        font_size=36,
        bold=True,
    )
    await add_heading(
        str(document_path),
        "PROFESSIONAL SUMMARY",
        level=2,
        font_name="Helvetica",
        font_size=14,
        bold=True,
        border_bottom=True,
    )
    await add_paragraph(
        str(document_path),
        "This text is italic and blue.",
        font_name="Arial",
        font_size=12,
        italic=True,
        color="0000FF",
    )

    document = Document(str(document_path))
    paragraphs = {paragraph.text: paragraph for paragraph in document.paragraphs}

    name_run = paragraphs["JAMES MEHORTER"].runs[0]
    assert name_run.font.name == "Helvetica"
    assert name_run.font.size is not None
    assert name_run.font.size.pt == 36
    assert name_run.bold is True

    summary_run = paragraphs["PROFESSIONAL SUMMARY"].runs[0]
    assert summary_run.font.name == "Helvetica"
    assert summary_run.font.size is not None
    assert summary_run.font.size.pt == 14
    assert summary_run.bold is True

    blue_run = paragraphs["This text is italic and blue."].runs[0]
    assert blue_run.font.name == "Arial"
    assert blue_run.font.size is not None
    assert blue_run.font.size.pt == 12
    assert blue_run.italic is True
    assert blue_run.font.color.rgb is not None
    assert str(blue_run.font.color.rgb) == "0000FF"
