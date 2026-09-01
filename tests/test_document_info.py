import json
from pathlib import Path

import pytest
from docx import Document

from word_mcp_codemode_live.tools.files import get_document_info


@pytest.mark.asyncio
async def test_file_document_info_does_not_report_sections_as_pages(tmp_path: Path) -> None:
    filename = tmp_path / "two-sections.docx"
    document = Document()
    document.add_paragraph("First section")
    document.add_section()
    document.add_paragraph("Second section")
    document.save(str(filename))

    result = json.loads(await get_document_info(str(filename)))

    assert result["section_count"] == 2
    assert "page_count" not in result
