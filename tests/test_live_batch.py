import gc
import json
import sys
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document as DocxDocument
from fastmcp import Client

from word_mcp_codemode_live.core import word_com
from word_mcp_codemode_live.main import create_server
from word_mcp_codemode_live.tools import batch as live_batch_tools
from word_mcp_codemode_live.tools.capture import RenderedPage


class _UndoRecord:
    IsRecordingCustomRecord = False

    def __init__(self, control: SimpleNamespace) -> None:
        self.control = control
        self.started = 0
        self.ended = 0

    def StartCustomRecord(self, _name: str) -> None:
        self.started += 1
        self.IsRecordingCustomRecord = True
        self.control.List = lambda _index: _name

    def EndCustomRecord(self) -> None:
        self.ended += 1
        self.IsRecordingCustomRecord = False


class _Document:
    Name = "batch-test.docx"
    Saved = False

    def __init__(self) -> None:
        self.Content = SimpleNamespace(Text="before")
        self.Paragraphs = SimpleNamespace(Count=1)
        self.Tables = SimpleNamespace(Count=0)
        self.Sections = SimpleNamespace(Count=1)
        self.saved_count = 0
        self.undo_count = 0

    def ComputeStatistics(self, kind: int) -> int:
        return {0: 1, 2: 1}[kind]

    def Repaginate(self) -> None:
        pass

    def Save(self) -> None:
        self.saved_count += 1
        self.Saved = True

    def Undo(self, _times: int = 1) -> bool:
        self.undo_count += 1
        self.Content.Text = "before"
        return True

    def Activate(self) -> None:
        pass


class _App:
    ScreenUpdating = True

    def __init__(self, document: _Document) -> None:
        self.document = document
        self.ActiveDocument = document
        self.Options = SimpleNamespace(Pagination=True)
        control = SimpleNamespace(ListCount=1, List=lambda _index: "MCP: Batch Edit")
        self.CommandBars = SimpleNamespace(FindControl=lambda **_kwargs: control)
        self.UndoRecord = _UndoRecord(control)


@pytest.mark.asyncio
async def test_batch_edits_save_once_and_return_a_page(monkeypatch) -> None:
    document = _Document()
    app = _App(document)

    async def fake_edit(value: str, filename: str | None = None) -> str:
        assert filename == "batch-test.docx"
        document.Content.Text = value
        with word_com.undo_record(app, "inner"):
            pass
        return json.dumps({"success": True})

    monkeypatch.setattr(live_batch_tools, "_batch_tools", lambda: {"fake_edit": fake_edit})
    monkeypatch.setattr(word_com, "get_word_app", lambda: app)
    monkeypatch.setattr(word_com, "find_document", lambda _app, _filename: document)
    monkeypatch.setattr(
        live_batch_tools,
        "render_word_pages",
        lambda *_args, **_kwargs: [RenderedPage(1, b"png", 10, 20)],
    )

    result = await live_batch_tools.word_live_edit_batch(
        [{"tool": "fake_edit", "arguments": {"value": "after"}}],
        filename="batch-test.docx",
        contains_text=["after"],
        absent_text=["before"],
        capture_pages=[1],
    )

    structured = result.structured_content
    assert structured is not None
    assert structured["success"] is True
    assert structured["captured_pages"] == [1]
    assert structured["timings"]["total_seconds"] >= 0
    assert document.saved_count == 1
    assert app.UndoRecord.started == 1
    assert app.UndoRecord.ended == 1


@pytest.mark.asyncio
async def test_batch_rolls_back_when_an_operation_fails(monkeypatch) -> None:
    document = _Document()
    app = _App(document)

    async def fake_edit(value: str, filename: str | None = None) -> str:
        document.Content.Text = value
        return json.dumps({"success": True})

    async def fake_error(filename: str | None = None) -> str:
        return json.dumps({"error": "runtime failure"})

    monkeypatch.setattr(
        live_batch_tools,
        "_batch_tools",
        lambda: {"fake_edit": fake_edit, "fake_error": fake_error},
    )
    monkeypatch.setattr(word_com, "get_word_app", lambda: app)
    monkeypatch.setattr(word_com, "find_document", lambda _app, _filename: document)

    result = await live_batch_tools.word_live_edit_batch(
        [
            {"tool": "fake_edit", "arguments": {"value": "partial"}},
            {"tool": "fake_error", "arguments": {}},
        ],
        filename="batch-test.docx",
    )

    assert result.is_error is True
    structured = result.structured_content
    assert structured is not None
    assert structured["rolled_back"] is True
    assert structured["failed_operation"] == 1
    assert document.Content.Text == "before"
    assert document.saved_count == 0
    assert document.undo_count == 1


@pytest.mark.asyncio
async def test_batch_rejects_unknown_tools_before_editing(monkeypatch) -> None:
    document = _Document()
    app = _App(document)

    async def fake_edit(value: str, filename: str | None = None) -> str:
        document.Content.Text = value
        return json.dumps({"success": True})

    monkeypatch.setattr(live_batch_tools, "_batch_tools", lambda: {"fake_edit": fake_edit})
    monkeypatch.setattr(word_com, "get_word_app", lambda: app)
    monkeypatch.setattr(word_com, "find_document", lambda _app, _filename: document)

    result = await live_batch_tools.word_live_edit_batch(
        [
            {"tool": "fake_edit", "arguments": {"value": "partial"}},
            {"tool": "not_allowed", "arguments": {}},
        ],
        filename="batch-test.docx",
    )

    assert result.is_error is True
    structured = result.structured_content
    assert structured is not None
    assert structured["rolled_back"] is False
    assert document.Content.Text == "before"
    assert document.undo_count == 0


@pytest.mark.asyncio
async def test_batch_rejects_string_boolean_before_editing(monkeypatch) -> None:
    document = _Document()
    app = _App(document)

    async def fake_edit(enabled: bool, filename: str | None = None) -> str:
        document.Content.Text = str(enabled)
        return json.dumps({"success": True})

    monkeypatch.setattr(live_batch_tools, "_batch_tools", lambda: {"fake_edit": fake_edit})
    monkeypatch.setattr(word_com, "get_word_app", lambda: app)
    monkeypatch.setattr(word_com, "find_document", lambda _app, _filename: document)

    result = await live_batch_tools.word_live_edit_batch(
        [{"tool": "fake_edit", "arguments": {"enabled": "false"}}],
        filename="batch-test.docx",
    )

    assert result.is_error is True
    structured = result.structured_content
    assert structured is not None
    assert "requires JSON boolean" in structured["error"]
    assert document.Content.Text == "before"
    assert document.undo_count == 0


@pytest.mark.asyncio
async def test_batch_rejects_nested_filename_before_editing(monkeypatch) -> None:
    document = _Document()
    app = _App(document)

    async def fake_edit(value: str, filename: str | None = None) -> str:
        document.Content.Text = value
        return json.dumps({"success": True})

    monkeypatch.setattr(live_batch_tools, "_batch_tools", lambda: {"fake_edit": fake_edit})
    monkeypatch.setattr(word_com, "get_word_app", lambda: app)
    monkeypatch.setattr(word_com, "find_document", lambda _app, _filename: document)

    result = await live_batch_tools.word_live_edit_batch(
        [
            {
                "tool": "fake_edit",
                "arguments": {"value": "wrong", "filename": "other.docx"},
            }
        ],
        filename="batch-test.docx",
    )

    assert result.is_error is True
    structured = result.structured_content
    assert structured is not None
    assert "batch-level filename" in structured["error"]
    assert document.Content.Text == "before"


@pytest.mark.skipif(sys.platform != "win32", reason="Microsoft Word COM requires Windows")
@pytest.mark.asyncio
async def test_batch_groups_real_word_edits_and_renders_page(tmp_path: Path) -> None:
    import win32com.client

    path = tmp_path / "live-batch.docx"
    source = DocxDocument()
    source.add_paragraph("Alpha placeholder")
    source.add_paragraph("Beta placeholder")
    source.save(str(path))

    try:
        app = win32com.client.DispatchEx("Word.Application")
    except Exception as exc:
        pytest.skip(f"Microsoft Word is unavailable: {exc}")

    app.Visible = False
    app.DisplayAlerts = 0
    document = None
    try:
        document = app.Documents.Open(str(path), ReadOnly=False, AddToRecentFiles=False)
        word_com.remember_word_app(app)

        async with Client(create_server(tool_mode="code")) as client:
            result = await client.call_tool(
                "word_live_edit_batch",
                {
                    "operations": [
                        {
                            "tool": "word_live_replace_text",
                            "arguments": {
                                "find_text": "Alpha placeholder",
                                "replace_text": "Alpha verified",
                            },
                        },
                        {
                            "tool": "word_live_replace_text",
                            "arguments": {
                                "find_text": "Beta placeholder",
                                "replace_text": "Beta verified",
                            },
                        },
                    ],
                    "filename": str(path),
                    "contains_text": ["Alpha verified", "Beta verified"],
                    "absent_text": ["Alpha placeholder", "Beta placeholder"],
                },
            )

        assert result.is_error is not True
        assert result.structured_content["captured_pages"] == [1]
        assert any(item.type == "image" for item in result.content)
        assert document.Undo(1)
        assert "Alpha placeholder" in document.Content.Text
        assert "Beta placeholder" in document.Content.Text
        assert "Alpha verified" not in document.Content.Text
        assert "Beta verified" not in document.Content.Text
    finally:
        if document is not None:
            document.Close(SaveChanges=False)
            document = None
        word_com._WORD_APP = None
        gc.collect()
        app.Quit(SaveChanges=False)
        app = None
